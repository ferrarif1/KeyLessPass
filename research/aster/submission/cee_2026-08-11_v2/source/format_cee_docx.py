#!/usr/bin/env python3
"""Apply a restrained Elsevier review-manuscript style to Pandoc DOCX files."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


TITLE = (
    "ASTER: Authorization-Scoped Threshold Exact-Domain Credential Derivation "
    "with Failure-Safe Root-Epoch Healing"
)


def set_font(run, name: str, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_layout_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, fractions, total_width=9120) -> None:
    widths = [round(total_width * fraction) for fraction in fractions]
    widths[-1] += total_width - sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill="F1F4F8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_minimal_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")
        borders.append(element)
    for edge in ("left", "right", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    set_font(run, "Times New Roman", 9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for child in (begin, instr, separate, value, end):
        run._r.append(child)


def suppress_line_numbers(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:suppressLineNumbers")) is None:
        p_pr.append(OxmlElement("w:suppressLineNumbers"))


def add_line_numbers(section) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:lnNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:distance"), "360")
    line_numbers.set(qn("w:restart"), "continuous")
    sect_pr.append(line_numbers)


def style_document(document: Document, mode: str) -> None:
    manuscript = mode == "manuscript"
    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(24)
        section.bottom_margin = Mm(22)
        section.left_margin = Mm(27 if manuscript else 25)
        section.right_margin = Mm(22 if manuscript else 25)
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)
        section.different_first_page_header_footer = False
        for paragraph in section.header.paragraphs:
            paragraph.text = ""
            suppress_line_numbers(paragraph)
        for paragraph in section.footer.paragraphs:
            paragraph.text = ""
            suppress_line_numbers(paragraph)
        if manuscript:
            add_page_number(section.footer.paragraphs[0])
            add_line_numbers(section)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading_tokens = {
        "Heading 1": (16, 14, 8, False),
        "Heading 2": (13, 12, 5, False),
        "Heading 3": (11.5, 9, 4, False),
        "Heading 4": (11, 7, 3, True),
    }
    for name, (size, before, after, italic) in heading_tokens.items():
        if name not in document.styles:
            continue
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.italic = italic
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    abstract_seen = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        has_drawing = bool(paragraph._p.xpath(".//w:drawing | .//w:pict"))
        has_display_math = bool(paragraph._p.xpath(".//m:oMathPara"))
        for run in paragraph.runs:
            set_font(run, "Times New Roman")

        if text == TITLE:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(12)
            for run in paragraph.runs:
                set_font(run, "Times New Roman", 16, True)
        elif manuscript and text == "Abstract":
            abstract_seen = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif manuscript and not abstract_seen and text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style_name.startswith("Heading"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif style_name in ("Source Code", "Verbatim Char"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            for run in paragraph.runs:
                set_font(run, "Liberation Mono", 8.5)
        elif has_display_math:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.keep_together = True
        elif has_drawing:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(7)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.keep_with_next = True
        elif mode == "title_page" and text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif mode == "highlights" and text == "Highlights":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif manuscript:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if text.startswith("Table "):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(7)
            paragraph.paragraph_format.space_after = Pt(4)
            # The compact candidate-descriptor table fits on one page but can
            # otherwise orphan its last row after the preceding figure. Start
            # it on a fresh page so the complete protocol object remains
            # visually atomic in the review manuscript.
            if text.startswith("Table 2."):
                paragraph.paragraph_format.page_break_before = True
            for run in paragraph.runs:
                set_font(run, "Times New Roman", 9)
        elif text.startswith("Figure "):
            paragraph.paragraph_format.keep_with_next = False
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(8)
            for run in paragraph.runs:
                set_font(run, "Times New Roman", 9)

    if manuscript:
        max_width = Mm(158)
        for shape in document.inline_shapes:
            if shape.width > max_width:
                ratio = max_width / shape.width
                shape.width = max_width
                shape.height = round(shape.height * ratio)

    table_widths = {
        "Approach": ([0.15, 0.10, 0.13, 0.12, 0.12, 0.12, 0.13, 0.13], 7.4),
        "Field": ([0.34, 0.66], 8.8),
        "Current state": ([0.24, 0.31, 0.45], 8.2),
        "Test": ([0.29, 0.49, 0.22], 8.6),
        "Rejection reason": ([0.78, 0.22], 8.8),
        "Metric": ([0.67, 0.33], 8.8),
        "Capability binding": ([0.36, 0.16, 0.16, 0.16, 0.16], 8.4),
        "Evidence": ([0.74, 0.26], 8.6),
        "Parties": ([0.10, 0.14, 0.10, 0.17, 0.13, 0.12, 0.12, 0.12], 7.4),
        "Rows": ([0.14, 0.27, 0.29, 0.30], 8.3),
    }
    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_layout_fixed(table)
        set_minimal_borders(table)
        first_header = table.rows[0].cells[0].text.strip() if table.rows else ""
        fractions, table_font_size = table_widths.get(
            first_header,
            ([1 / len(table.columns)] * len(table.columns), 8.6),
        )
        set_table_widths(table, fractions)
        if table.rows:
            set_repeat_header(table.rows[0])
        for row_index, row in enumerate(table.rows):
            set_row_cant_split(row)
            for col_index, cell in enumerate(row.cells):
                set_cell_margins(cell)
                if row_index == 0:
                    shade_cell(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.05
                    if row_index == 0 or col_index > 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        set_font(run, "Times New Roman", table_font_size, row_index == 0)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument(
        "--mode",
        choices=("manuscript", "title_page", "highlights", "letter", "declarations", "checklist"),
        required=True,
    )
    args = parser.parse_args()
    document = Document(args.input)
    style_document(document, args.mode)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
