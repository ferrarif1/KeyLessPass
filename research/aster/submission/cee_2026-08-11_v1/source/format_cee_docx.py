#!/usr/bin/env python3
"""Apply a restrained Elsevier review-manuscript style to Pandoc DOCX files."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


TITLE = (
    "ASTER: Authorization-Scoped Threshold Exact-Domain Credential Derivation "
    "with Failure-Safe Root-Epoch Healing"
)


def set_font(run, name: str, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_layout_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_minimal_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")
        borders.append(element)
    for edge in ("left", "right", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    tbl_pr.append(borders)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    set_font(run, "Times New Roman", 9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for child in (begin, instr, separate, value, end):
        run._r.append(child)


def suppress_line_numbers(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:suppressLineNumbers")) is None:
        p_pr.append(OxmlElement("w:suppressLineNumbers"))


def add_line_numbers(section) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:lnNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:distance"), "360")
    line_numbers.set(qn("w:restart"), "continuous")
    sect_pr.append(line_numbers)


def style_document(document: Document, mode: str) -> None:
    manuscript = mode == "manuscript"
    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(24)
        section.bottom_margin = Mm(22)
        section.left_margin = Mm(27 if manuscript else 25)
        section.right_margin = Mm(22 if manuscript else 25)
        section.header_distance = Mm(10)
        section.footer_distance = Mm(10)
        section.different_first_page_header_footer = False
        for paragraph in section.header.paragraphs:
            paragraph.text = ""
            suppress_line_numbers(paragraph)
        for paragraph in section.footer.paragraphs:
            paragraph.text = ""
            suppress_line_numbers(paragraph)
        if manuscript:
            add_page_number(section.footer.paragraphs[0])
            add_line_numbers(section)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heading_tokens = {
        "Heading 1": (16, 14, 8, False),
        "Heading 2": (13, 12, 5, False),
        "Heading 3": (11.5, 9, 4, False),
        "Heading 4": (11, 7, 3, True),
    }
    for name, (size, before, after, italic) in heading_tokens.items():
        if name not in document.styles:
            continue
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.italic = italic
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    abstract_seen = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        for run in paragraph.runs:
            set_font(run, "Times New Roman")

        if text == TITLE:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(12)
            for run in paragraph.runs:
                set_font(run, "Times New Roman", 16, True)
        elif manuscript and text == "Abstract":
            abstract_seen = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif manuscript and not abstract_seen and text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif style_name.startswith("Heading"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif mode == "title_page" and text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif mode == "highlights" and text == "Highlights":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif manuscript:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if text.startswith(("Figure ", "Table ")):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(7)
            paragraph.paragraph_format.space_after = Pt(4)
            for run in paragraph.runs:
                set_font(run, "Times New Roman", 9)

    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        set_table_layout_fixed(table)
        set_minimal_borders(table)
        if table.rows:
            set_repeat_header(table.rows[0])
        for row_index, row in enumerate(table.rows):
            set_row_cant_split(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        set_font(run, "Times New Roman", 9, row_index == 0)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument(
        "--mode",
        choices=("manuscript", "title_page", "highlights", "letter", "declarations", "checklist"),
        required=True,
    )
    args = parser.parse_args()
    document = Document(args.input)
    style_document(document, args.mode)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
