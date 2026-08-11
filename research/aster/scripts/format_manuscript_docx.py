#!/usr/bin/env python3
"""Apply final ASTER manuscript formatting after Pandoc conversion."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def replace_story_text(paragraphs, text: str) -> None:
    for paragraph in paragraphs:
        paragraph.text = text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    document = Document(args.input)

    for section in document.sections:
        replace_story_text(section.header.paragraphs, "ASTER")
        replace_story_text(section.first_page_header.paragraphs, "")
        replace_story_text(section.even_page_header.paragraphs, "ASTER")
        replace_story_text(section.footer.paragraphs, "")
        replace_story_text(section.first_page_footer.paragraphs, "")
        replace_story_text(section.even_page_footer.paragraphs, "")

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Yuanyi Zhang\nHangzhou Information Technology Branch"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text in {
            "ASTER: Authorization-Scoped Threshold Exact-Domain Credential Derivation with Failure-Safe Root-Epoch Healing",
        }:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text.startswith("Table 6. Fault-injection"):
            paragraph.paragraph_format.page_break_before = True

    for table in document.tables:
        if table.rows:
            set_repeat_header(table.rows[0])
        for row in table.rows:
            set_row_cant_split(row)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
