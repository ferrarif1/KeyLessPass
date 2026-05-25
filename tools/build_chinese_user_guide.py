#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "KeyLessPass_中文使用手册.docx"
ASSETS = ROOT / "docs" / "user-guide-assets" / "zh-screenshots"
LOGO = ROOT / "docs" / "readme-assets" / "logo.png"
BODY_FONT = "PingFang SC"


COLORS = {
    "ink": RGBColor(20, 24, 31),
    "muted": RGBColor(92, 99, 112),
    "blue": RGBColor(46, 116, 181),
    "dark_blue": RGBColor(31, 77, 120),
    "primary": RGBColor(246, 255, 84),
    "border": "DADCE0",
    "header_fill": "E8EEF5",
    "callout_fill": "F4F6F9",
    "warning_fill": "FFF8E1",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "DADCE0", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_run_font(run, name: str = BODY_FONT, size: float | None = None, bold: bool | None = None,
                 color: RGBColor | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc: Document, text: str = "", style: str | None = None, *,
                  bold_prefix: str | None = None) -> object:
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run_font(run)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level], bold=True,
                 color=COLORS["blue"] if level < 3 else COLORS["dark_blue"])


def add_callout(doc: Document, title: str, body: str, fill: str = "F4F6F9") -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="DADCE0")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, bold=True, color=COLORS["dark_blue"])
    p.add_run("\n")
    r2 = p.add_run(body)
    set_run_font(r2, size=10.5)
    doc.add_paragraph()


def add_label_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table)
    widths = [Inches(1.6), Inches(4.9)]
    for row_idx, (label, value) in enumerate(rows):
        cells = table.rows[row_idx].cells
        for i, cell in enumerate(cells):
            cell.width = widths[i]
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], COLORS["header_fill"])
        p0 = cells[0].paragraphs[0]
        r0 = p0.add_run(label)
        set_run_font(r0, bold=True, color=COLORS["dark_blue"])
        p1 = cells[1].paragraphs[0]
        r1 = p1.add_run(value)
        set_run_font(r1)
    doc.add_paragraph()


def add_screenshot(doc: Document, image_name: str, caption: str) -> None:
    path = ASSETS / image_name
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.45))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, size=9.5, color=COLORS["muted"])


def prepare_redacted_screenshots() -> None:
    """Create documentation-safe derivatives for screenshots that show secrets."""
    source = ASSETS / "口令派生5.png"
    target = ASSETS / "口令派生5_密码已遮盖.png"
    if not source.exists():
        return
    image = Image.open(source).convert("RGB")
    draw = ImageDraw.Draw(image)
    # Cover the generated password field while keeping the success state visible.
    draw.rectangle((690, 910, 1350, 1048), fill=(20, 24, 31), outline=(246, 255, 84), width=3)
    draw.text((745, 955), "密码已遮盖", fill=(246, 255, 84))
    image.save(target)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLORS["ink"]
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, left_indent in (("List Bullet", 0.38), ("List Number", 0.38)):
        style = doc.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(left_indent)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build() -> None:
    prepare_redacted_screenshots()
    doc = Document()
    configure_styles(doc)

    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run().add_picture(str(LOGO), width=Inches(0.82))

    title = doc.add_paragraph()
    r = title.add_run("KeyLessPass 快速使用手册")
    set_run_font(r, size=24, bold=True, color=COLORS["ink"])
    subtitle = doc.add_paragraph()
    r = subtitle.add_run("给普通用户看的图文教程")
    set_run_font(r, size=13, color=COLORS["muted"])

    add_label_table(doc, [
        ("适用对象", "需要登录内部系统、运维平台、厂商系统的普通用户和管理员"),
        ("文档日期", "2026-05-25"),
        ("你需要准备", "一台已安装 KeyLessPass 的电脑、一块 U 盘、一组助记短语"),
        ("一句话说明", "KeyLessPass 不保存你的系统密码，只在你需要登录时临时生成。"),
    ])

    add_callout(
        doc,
        "先记住这三件事",
        "1. 助记短语不要丢，也不要发给别人。\n2. U 盘要妥善保管，建议专盘专用。\n3. 生成出来的密码只用于当前系统登录，不会被 KeyLessPass 保存。",
        "FFF8E1",
    )

    doc.add_page_break()

    add_heading(doc, "1. 第一次使用：完成初始化", 1)
    add_paragraph(doc, "初始化只需要做一次。请先插入 U 盘，然后按页面提示操作。")
    add_number(doc, "在左侧点击“初始化”。")
    add_number(doc, "选择“英文”或“简体中文”助记短语。")
    add_number(doc, "点击“生成助记短语”，把它抄写到安全的地方。不要截图，不要上传网盘。")
    add_number(doc, "确认 U 盘路径，例如 /Volumes/WD。如果没有自动识别，点击文件夹图标选择 U 盘。")
    add_number(doc, "点击“创建因子”。看到状态变成“已初始化”后，就可以添加记录。")
    add_screenshot(doc, "初始设置1.png", "图 1 初始化：生成助记短语，选择 U 盘，然后创建因子")

    add_heading(doc, "2. 添加一个系统记录", 1)
    add_paragraph(doc, "一个“记录”对应一个需要登录的系统。这里只保存系统名称、账号提示和密码规则，不保存真实密码。")
    add_number(doc, "点击左侧“记录”，再点击右上角“添加记录”。")
    add_number(doc, "填写显示名称，例如“运维控制台”。")
    add_number(doc, "填写服务提示或 URL，方便以后查找。")
    add_number(doc, "填写账号提示，例如 admin01。")
    add_number(doc, "选择密码长度和字符规则，然后点击“创建”。")
    add_callout(
        doc,
        "放心修改",
        "以后修改系统名称、URL 或账号提示，不会改变已经生成过的密码。只有你主动做“轮换”时，密码才会变。",
        "F4F6F9",
    )
    add_screenshot(doc, "口令派生2.png", "图 2 添加记录：填写系统名称、账号提示和密码规则")
    add_screenshot(doc, "口令派生3.png", "图 3 记录列表：查看、搜索和选择已有记录")

    add_heading(doc, "3. 生成登录密码", 1)
    add_paragraph(doc, "每次需要登录系统时，进入“派生密码”页面临时生成密码。生成后请尽快复制到目标系统使用。")
    add_number(doc, "在“派生密码”页面选择要登录的系统记录。")
    add_number(doc, "保持“本机验证”即可。只有换电脑或本机数据丢失时，才使用“U 盘恢复验证”。")
    add_number(doc, "输入助记短语。")
    add_number(doc, "确认 U 盘路径。")
    add_number(doc, "点击“派生并复制”。密码会复制到剪贴板，并会自动清除。")
    add_callout(
        doc,
        "为什么有 U 盘还要输入助记短语？",
        "因为 U 盘只是普通存储设备。KeyLessPass 需要“这台电脑 + U 盘 + 助记短语”一起确认，才会生成密码。",
        "FFF8E1",
    )
    add_screenshot(doc, "口令派生4.png", "图 4 派生密码：选择记录，输入助记短语和 U 盘路径")
    add_screenshot(doc, "口令派生5_密码已遮盖.png", "图 5 派生成功：示例密码已在本文档中遮盖")

    add_heading(doc, "4. 修改目标系统密码", 1)
    add_paragraph(doc, "当目标系统要求改密，或者你想主动更换密码时，使用“轮换”。")
    add_number(doc, "进入“轮换”，选择要改密的记录。")
    add_number(doc, "点击“创建待确认版本”。")
    add_number(doc, "生成新密码，把目标系统里的密码改成这个新密码。")
    add_number(doc, "确认目标系统改密成功后，再点击“提交轮换”。")
    add_number(doc, "如果目标系统改密失败，点击取消；旧密码仍然有效。")
    add_screenshot(doc, "轮换1.png", "图 6 轮换：选择记录并开始创建新版本")
    add_screenshot(doc, "轮换4.png", "图 7 轮换：生成新密码后，先去目标系统完成改密")

    add_heading(doc, "5. U 盘和记录备份", 1)
    add_paragraph(doc, "U 盘页面用于检查 U 盘是否可用，也用于同步记录备份。这里的备份不包含真实密码。")
    add_bullet(doc, "如果提示“本机和 U 盘不一致”，先判断哪一份是最新的。")
    add_bullet(doc, "如果本机记录是最新的，选择同步到 U 盘。")
    add_bullet(doc, "如果换电脑或本机数据丢失，可以选择从 U 盘恢复本机记录。")
    add_screenshot(doc, "记录备份同步.png", "图 8 U 盘设备：检查记录备份是否一致")

    add_heading(doc, "6. 丢失或更换设备时怎么办", 1)
    add_label_table(doc, [
        ("U 盘丢了", "准备原电脑和助记短语，进入恢复页面，选择恢复 U 盘。"),
        ("换电脑了", "准备 U 盘和助记短语，在新电脑上恢复本机材料。"),
        ("想换助记短语", "准备当前电脑和 U 盘，选择重置助记短语。旧助记短语会失效。"),
        ("要重新开始", "在设置里执行重置应用数据。此操作很危险，确认前请备份好 U 盘。"),
    ])
    add_screenshot(doc, "恢复u盘.png", "图 9 恢复 U 盘：U 盘丢失后重新创建 U 盘数据")
    add_screenshot(doc, "恢复本机.png", "图 10 恢复本机：换电脑时恢复本机数据")
    add_screenshot(doc, "重置助记短语.png", "图 11 重置助记短语：用当前电脑和 U 盘设置新助记短语")

    add_heading(doc, "7. 常见问题", 1)
    faq = [
        ("KeyLessPass 会保存我的系统密码吗？", "不会。密码只在你点击“派生并复制”时临时生成。"),
        ("助记短语会保存在哪里？", "不会保存。请自己离线保存。"),
        ("改了系统名称或 URL，密码会变吗？", "不会。名称和 URL 只是方便查找。"),
        ("什么时候密码会变？", "只有进行“轮换”并提交后，后续生成的密码才会变。"),
        ("U 盘没有识别怎么办？", "在 Finder 中确认 U 盘存在，然后点击界面里的文件夹图标选择 U 盘目录。"),
    ]
    add_label_table(doc, faq)

    add_heading(doc, "8. 日常安全建议", 1)
    add_bullet(doc, "助记短语要离线保存，不要拍照，不要发聊天工具。")
    add_bullet(doc, "U 盘建议专盘专用，平时妥善保管。")
    add_bullet(doc, "生成密码后尽快使用，离开页面前确认密码已隐藏或清除。")
    add_bullet(doc, "轮换密码前，先确认目标系统允许你修改密码。")
    add_bullet(doc, "不要把包含真实系统名、账号或 IP 的截图发给无关人员。")

    # Footer
    for section in doc.sections:
        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run("KeyLessPass 本机密码派生客户端 · 中文使用手册")
        set_run_font(run, size=9, color=COLORS["muted"])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
